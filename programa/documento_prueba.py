from docx import Document
from sympy import sympify, Matrix

def extraer_ecuaciones_y_matrices(docx_file):
    ecuaciones = []
    matrices = []

    document = Document(docx_file)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        # Buscar ecuaciones entre corchetes [ ]
        if "[" in text and "]" in text:
            ecuacion = text[text.find("[")+1:text.find("]")]
            try:
                ecuacion_sympy = sympify(ecuacion)
                ecuaciones.append(ecuacion_sympy)
            except:
                pass

        # Buscar matrices
        if "MATRIZ" in text:
            table = paragraph.tables[0]  # Suponiendo que la matriz está en la primera tabla
            valores = []
            for row in table.rows:
                fila = []
                for cell in row.cells:
                    fila.append(cell.text.strip())
                valores.append(fila)
            
            matriz_sympy = Matrix(valores)
            matrices.append(matriz_sympy)

    return ecuaciones, matrices
# Ejemplo de uso
documento = "./programa//doc/matriz1.docx"
ecuaciones_extraidas, matrices_extraidas = extraer_ecuaciones_y_matrices(documento)

print(matrices_extraidas)
print(ecuaciones_extraidas)